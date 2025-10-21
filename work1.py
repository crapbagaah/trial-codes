from docx import Document
import os

def extract_images_from_run(run, image_folder, image_map, image_counter):
    """
    Extracts images from a run, saves them to disk, and returns Markdown link text.
    """
    md_text = ""
    drawing_elems = run._element.findall(".//a:blip", run._element.nsmap)
    for d in drawing_elems:
        rid = d.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rid not in image_map:
            image_part = run.part.related_parts[rid]
            image_data = image_part.blob
            img_name = f"image_{image_counter[0]}.png"
            img_path = os.path.join(image_folder, img_name)
            with open(img_path, "wb") as f:
                f.write(image_data)
            image_map[rid] = img_name
            image_counter[0] += 1
        md_text += f"\n\n![{image_map[rid]}](images/{image_map[rid]})\n\n"
    return md_text


def split_docx_to_markdown(docx_path, output_folder, expected_pages=None):
    doc = Document(docx_path)
    os.makedirs(output_folder, exist_ok=True)
    image_folder = os.path.join(output_folder, "images")
    os.makedirs(image_folder, exist_ok=True)

    image_map = {}
    image_counter = [1]

    pages = []
    current_page = ""

    def save_page():
        nonlocal current_page
        if current_page.strip():
            pages.append(current_page.strip())
            current_page = ""

    for block in doc.element.body:
        # Handle paragraphs
        if block.tag.endswith("p"):
            for para in doc.paragraphs:
                if para._element == block:
                    xml = para._element.xml

                    # Detect manual or section page break
                    if '<w:br w:type="page"' in xml or 'w:type="nextPage"' in xml:
                        save_page()
                        break

                    # Text and image runs
                    for run in para.runs:
                        if run.text.strip():
                            current_page += run.text + " "

                        if "graphic" in run._element.xml:
                            current_page += extract_images_from_run(run, image_folder, image_map, image_counter)

                    current_page += "\n\n"
                    break

        # Handle tables
        elif block.tag.endswith("tbl"):
            rows = block.findall(".//w:tr")
            for row in rows:
                cells = row.findall(".//w:tc")
                cell_texts = [" ".join([t.text or "" for t in c.findall(".//w:t")]) for c in cells]
                current_page += "| " + " | ".join(cell_texts) + " |\n"
            current_page += "\n"

    save_page()  # last page

    # Fallback if no explicit breaks and user expects N pages
    if expected_pages and len(pages) < expected_pages:
        # Roughly divide by paragraph count
        total_paras = len(doc.paragraphs)
        split_size = total_paras // expected_pages
        pages = []
        for i in range(0, total_paras, split_size):
            chunk = doc.paragraphs[i:i+split_size]
            text_chunk = "\n".join([p.text for p in chunk])
            pages.append(text_chunk.strip())

    # Write Markdown files
    for i, content in enumerate(pages, start=1):
        with open(os.path.join(output_folder, f"page{i}.md"), "w", encoding="utf-8") as f:
            f.write(f"# Page {i}\n\n{content}\n")

    print(f"✅ Created {len(pages)} Markdown pages with inline images in {output_folder}/")


# Example usage:
# If your DOCX has exactly 8 pages, you can specify expected_pages=8
split_docx_to_markdown("input.docx", "output_md", expected_pages=8)
